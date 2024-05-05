from docx import Document
from docx.shared import Inches

def replace_text(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    cell.text = cell.text.replace(old_text, new_text)

def main():
    # Открываем шаблонный файл
    template_path = 'шаблон.docx'
    doc = Document(template_path)

    # Заменяем данные в шаблоне
    replace_text(doc, 'OLD_TEXT', 'Соси')
    # Можно добавить сколько угодно таких замен

    # Сохраняем измененный файл
    doc.save('output.docx')

if __name__ == "__main__":
    main()
# import streamlit as st
# from docx import Document
# import datetime
# from utils import replace_text

# def main():
#     # Открываем шаблонный файл
#     template_path = 'template.docx'
#     doc = Document(template_path)

#     # Словарь для замены значений в шаблоне
#     replacements = {}

#      # Получаем значения из Streamlit input
#     kn_input_value = st.text_input("Введите Кадастровый", "")
#     replacements['<КадастровыйНомер>'] = kn_input_value

#     # Получаем значения из Streamlit input
#     fio_input_value = st.text_input("Введите ФИО п.1.3", "")
#     replacements['<ФИО>'] = fio_input_value

#     # Получаем значения из Streamlit input
#     num_request_input_value = st.text_input("Номер завявления", "")
#     replacements['<НомерЗаявления>'] = num_request_input_value
   
#     request_date_value = st.date_input("Дата поступления обращения","today",format='DD/MM/YYYY')
#     replacements['<ДатаПоступленияЗаявления>'] = str(request_date_value.strftime('%d.%m.%Y'))

#     email = st.text_input("Почта")
#     replacements['<Почта>'] = email

#     # Кнопка для замены данных в шаблоне
#     if st.button("Заменить данные в шаблоне"):
#         replace_text(doc, replacements)
#         st.success("Данные успешно заменены в шаблоне!")

#         # Сохраняем измененный файл
#         doc.save('output.docx')

# if __name__ == "__main__":
#     main()


    # def tables_page1(self):

    #     # Check if 'Кадастровый номер' exists in the data columns
    #     if 'Кадастровый номер' in self.data.columns:
    #         # Find rows where 'Кадастровый номер' matches kn_input_value
    #         matches = self.data.loc[self.data['Кадастровый номер'] == self.replacements.conditions['<КадастровыйНомер>'], 
    #                                 ['Наименование здания', 'Назначение/Проектируемое назначение здания', 'Вид объекта недвижимости','Площадь ЕГРН в 1 лист','Адрес','Кол-во надземных этажей ЕГРН','Кол-во подземных этажей ЕГРН',
    #                                  'Материал стен ЕГРН','Год ввода в эксплуатацию ЕГРН','Площадь, иная характеристика в 5.3','Кадастровый номер земельного участка','кол-во надземных этажей в 5.3','Количество подземных этажей в 5.3',
    #                                  'Материал стен в 5.3','Год ввода в эксплуатацию в 5.3']]

    #         # Check if there are any matches
    #         if not matches.empty:
    #             # таблица 1
    #             building_name = matches.iloc[0, 0]  # First column
    #             purpose = matches.iloc[0, 1]  # Second column
    #             type_p = matches.iloc[0, 2]  # Third column
    #             area_egrn = str(matches.iloc[0, 3])  # Third column
    #             adress = matches.iloc[0, 4]  # Third column
    #             postive_floor = str(matches.iloc[0, 5])  # Third column
    #             negative_floor = str(matches.iloc[0, 6])  # Third column
    #             wall_type = matches.iloc[0, 7]  # Third column
    #             foundation_year = str(matches.iloc[0, 8])  # Third column
    #             kn_zu = matches.iloc[0, 10]
    #             if kn_zu == 'nan':
    #                 kn_zu = str(kn_zu)  # Third column
    #             else: kn_zu = '-'
    #             #таблица 2
    #             area2 = str(matches.iloc[0, 9])  # Third column
    #             postive_floor2 = str(matches.iloc[0, 11])  # Third column
    #             negative_floor2 = str(matches.iloc[0, 12])  # Third column
    #             wall_type2 = str(matches.iloc[0, 13])
    #             foundation_year2 = str(matches.iloc[0, 14])

    #             # Update the conditions
    #             # таблица 1
    #             self.replacements.update_conditions('<Наименование>', anchor=building_name)
    #             self.replacements.update_conditions('<Назначение>', anchor=purpose)
    #             self.replacements.update_conditions('<Вид_объекта>', anchor=type_p)
    #             self.replacements.update_conditions('<Площадь>', anchor=area_egrn)
    #             self.replacements.update_conditions('<Адрес>', anchor=adress)
    #             self.replacements.update_conditions('<Этажи+>', anchor=postive_floor)
    #             self.replacements.update_conditions('<Этажи->', anchor=negative_floor)
    #             self.replacements.update_conditions('<Материал_стен>', anchor=wall_type)
    #             self.replacements.update_conditions('<Год_ввода>', anchor=foundation_year)
    #             self.replacements.update_conditions('<КН_ЗУ>', anchor=kn_zu)

    #             #таблица 2
    #             self.replacements.update_conditions('<2Этажи+>', anchor=postive_floor2)
    #             self.replacements.update_conditions('<2Этажи->', anchor=negative_floor2)
    #             self.replacements.update_conditions('<2Материал_стен>', anchor=wall_type2)
    #             self.replacements.update_conditions('<2Год_ввода>', anchor=foundation_year2)
    #             self.replacements.update_conditions('<2Площадь>', anchor=area2)



    #         else:
    #             print(f"No building name found for 'Кадастровый номер': {self.replacements.conditions['<КадастровыйНомер>']}")
    #     else:
    #         print("'Кадастровый номер' column not found in the Excel data.")
