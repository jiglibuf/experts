from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def get_text_style(paragraph):
    style = {}
    style['text'] = paragraph.text
    style['bold'] = paragraph.runs[0].bold
    style['italic'] = paragraph.runs[0].italic
    style['underline'] = paragraph.runs[0].underline
    style['font'] = paragraph.runs[0].font
    style['alignment'] = paragraph.alignment
    return style

def apply_text_style(new_run, font):
    new_run.bold = font.bold
    new_run.italic = font.italic
    new_run.underline = font.underline
    new_run.font.name = font.name
    new_run.font.size = font.size
    new_run.font.color.rgb = font.color.rgb

def replace_text(doc, replacements):
    def replace_anchors(text):
        # Шаблон для поиска якорей
        pattern = r'<[^>]*>'

        # Заменяем все найденные якоря на их значения
        replaced_text = re.sub(pattern, lambda m: replacements.get(m.group(0), ''), text)

        return replaced_text

    # Заменяем текст в колонтитулах
    for section in doc.sections:
        if section.header is not None:
            for paragraph in section.header.paragraphs:
                # Проверяем наличие якорей в тексте параграфа
                if any(anchor in paragraph.text for anchor in replacements.keys()):
                    new_text = replace_anchors(paragraph.text)
                    for run in paragraph.runs:
                        run.clear()
                    new_run = paragraph.add_run(new_text)
                    apply_text_style(new_run, paragraph.runs[0].font)

        if section.footer is not None:
            for paragraph in section.footer.paragraphs:
                # Проверяем наличие якорей в тексте параграфа
                if any(anchor in paragraph.text for anchor in replacements.keys()):
                    new_text = replace_anchors(paragraph.text)
                    for run in paragraph.runs:
                        run.clear()
                    new_run = paragraph.add_run(new_text)
                    apply_text_style(new_run, paragraph.runs[0].font)

    # Заменяем текст в основном тексте документа
    for paragraph in doc.paragraphs:
        # Проверяем наличие якорей в тексте параграфа
        if any(anchor in paragraph.text for anchor in replacements.keys()):
            new_text = replace_anchors(paragraph.text)
            for run in paragraph.runs:
                run.clear()
            new_run = paragraph.add_run(new_text)
            apply_text_style(new_run, paragraph.runs[0].font)

    # Заменяем текст в якорях таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    # Проверяем наличие якорей в тексте параграфа
                    if any(anchor in paragraph.text for anchor in replacements.keys()):
                        new_text = replace_anchors(paragraph.text)
                        for run in paragraph.runs:
                            run.clear()
                        new_run = paragraph.add_run(new_text)
                        apply_text_style(new_run, paragraph.runs[0].font)

    # Удаление пустых строк
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == '':
            paragraph.clear()


def update_conditions(replacements, condition_to_update, anchor=None):
    if anchor is None:
        anchor = 'V'

    # Добавляем ключ, если его нет
    if condition_to_update not in replacements:
        replacements[condition_to_update] = anchor
    else:
        # Если ключ уже есть, обновляем значение
        replacements[condition_to_update] = anchor

    # Возвращаем обновленный словарь с условиями
    return replacements

